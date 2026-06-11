import logging
from io import BytesIO
from pathlib import Path
from typing import List

from docx import Document

logger = logging.getLogger(__name__)

OUTPUT_PATH = "app/output"


def table_to_markdown(table):
    """Converts a docx table to a markdown table."""
    md = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.replace("\n", " ") for cell in row.cells]
        line = "| " + " | ".join(cells) + " |"
        md.append(line)
        if i == 0:
            md.append("|" + "---|" * len(cells))
    return "\n".join(md)


def extract_fragments(file_content: bytes):
    """
    Extracts text and tables from a docx file and returns them as a list of
     markdown fragments.
    """
    doc = Document(BytesIO(file_content))
    fragments = []
    current = ""
    tables = list(doc.tables)

    for block in doc.element.body:
        if block.tag.endswith("tbl"):
            table = tables.pop(0)
            md_table = table_to_markdown(table)
            if len(current) + len(md_table) + 2 <= 1000:
                current += ("\n" if current else "") + md_table
            else:
                if current:
                    fragments.append(current)
                current = md_table
        elif block.tag.endswith("p"):
            p = block
            # Detect paragraph style for markdown heading
            para_obj = None
            for para in doc.paragraphs:
                if para._element == p:
                    para_obj = para
                    break
            if para_obj:
                style = para_obj.style.name
                paragraph = para_obj.text
                if style.startswith("Heading"):
                    # Extract heading level
                    try:
                        level = int(style.replace("Heading ", ""))
                    except Exception:
                        level = 1
                    heading_md = "#" * level + " " + paragraph
                    paragraph = heading_md
            else:
                paragraph = p.text
            if not paragraph.strip():
                continue
            if len(current) + len(paragraph) + 2 <= 1000:
                current += ("\n" if current else "") + paragraph
            else:
                if current:
                    fragments.append(current)
                current = paragraph
    if current:
        fragments.append(current)
    return fragments


def write_fragments_to_md(fragments, filename):
    """Writes a list of fragments to a markdown file."""
    md_name = filename.replace(".docx", ".md")
    with open(md_name, "w", encoding="utf-8") as f:
        for fragment in fragments:
            f.write(f"\n{fragment}\n---\n")


def merge_fragments(fragments: List[str]) -> str:
    """Merges a list of fragments into a single string."""
    full_text = []
    for frag in fragments:
        full_text.append(frag)
    return "\n".join(full_text)


def extract_text_from_docx(file_content: bytes, filename: str) -> str:
    """
    Extracts text from a docx file, saves it as a markdown file, and returns
    the full text.
    """
    fragments = extract_fragments(file_content)
    write_fragments_to_md(fragments, filename)
    full_text = []
    for frag in fragments:
        full_text.append(frag)
    return "\n".join(full_text)


def extract_images(file_content: bytes, filename: str):
    """Extracts images from a docx file and saves them to the output directory."""
    name = filename
    output_dir = Path(f"app/output/{name}")
    output_dir.mkdir(exist_ok=True)
    (output_dir / "images").mkdir(parents=True, exist_ok=True)
    doc = Document(BytesIO(file_content))

    images = []
    image_count = 0

    drawings = doc.element.body.xpath(".//a:blip")
    for drawing_index, drawing in enumerate(drawings, start=1):
        rId = drawing.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )
        if not rId:
            continue

        image_part = doc.part.related_parts.get(rId)
        if image_part is None:
            logger.warning(
                "Unable to resolve image relationship %s for drawing %s",
                rId,
                drawing_index,
            )
            continue

        image_bytes = image_part.blob
        image_count += 1
        ext = image_part.partname.ext.replace(".", "") or "bin"
        img_path = output_dir / f"images/{name}-picture-{image_count:03d}.{ext}"
        logger.info("Writing image %s to %s", drawing_index, img_path)
        with open(img_path, "wb") as f:
            f.write(image_bytes)
        images.append(img_path)

    return images


def process_docx_file(file_content: bytes, filename: str):
    """Processes a docx file and returns the full text, images, tables, and pages."""
    full_text = extract_text_from_docx(file_content, filename)
    images = extract_images(file_content, filename)
    tables = []  # Placeholder for table extraction logic
    pages = []  # Placeholder for page extraction logic
    return full_text, images, tables, pages
